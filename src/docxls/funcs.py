import os
import csv
import pandas as pd
from docx import Document
import subprocess
from config import get_logger
logger = get_logger(__name__)


def convert_to_new_format(file_path):
    """Convert .doc -> .docx and .xls -> .xlsx using LibreOffice."""
    ext = os.path.splitext(file_path)[1].lower()
    output_dir = os.path.dirname(file_path)
    if ext in [".doc", ".xls"]:
        logger.info(f"Converting {file_path} ...")
        subprocess.run([
            "libreoffice", "--headless", "--convert-to",
            "docx" if ext == ".doc" else "xlsx",
            file_path, "--outdir", output_dir
        ], check=True)
        new_file = file_path + "x"
        return new_file
    return file_path

def read_docx(path):
    doc = Document(path)
    contents = []
    text_lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            text_lines.append(p.text.strip())
    if len(text_lines) > 0:
        contents.append({"text": text_lines})

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_data.append(row_text)
        contents.append({"cells": table_data})
    
    return contents


def read_txt(path):
    with open(path, encoding="utf-8", errors="ignore") as f:
        lines = [line.strip() for line in f if line.strip()]
    return [{"text": lines}]


def read_csv(path):
    text_lines = []
    table_cells = []
    with open(path, encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            if len(row) == 1:
                text_lines.append(row[0])
            else:
                table_cells.append(row)
    return [{"text": text_lines, "cells": table_cells}]


def read_xlsx(path):
    pages = []
    try:
        xls = pd.ExcelFile(path)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
            df = df.dropna(how="all")
            if not df.empty:
                cells = df.fillna("").astype(str).values.tolist()
                pages.append({"cells": cells, "sheet": sheet_name})
    except Exception as e:
        print(f"⚠️ Error reading {path}: {e}")
    return pages


def read_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    file_path = convert_to_new_format(file_path)
    pages = []
    if ext in [".doc", ".docx"]:
        pages = read_docx(file_path)
    elif ext == ".txt":
        pages = read_txt(file_path)
    elif ext in [".xls", ".xlsx"]:
        pages = read_xlsx(file_path)
    elif ext == ".csv":
        pages = read_csv(file_path)
    else:
        pages = [{"text": ["Unsupported format"]}]
    return pages

def make_avatar_file(file_path):
    basename = os.path.basename(file_path)
    filename, _ = os.path.splitext(basename)
    output_dir = os.path.dirname(file_path)

    subprocess.run([
        "libreoffice",
        "--headless",
        "--convert-to", "png",
        "--outdir", output_dir,
        file_path
    ], check=True)
    
    filename = os.path.splitext(os.path.basename(file_path))[0]
    png_path = os.path.join(output_dir, f"{filename}.png")
    logger.info(f"✅ Avatar created for file {basename}: {png_path}")
    return png_path